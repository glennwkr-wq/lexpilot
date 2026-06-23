from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def build_legal_docx(
    content: str,
    title: str = "Юридический документ",
    client_name: str = "",
) -> BytesIO:
    document = Document()

    _setup_page(document)
    _setup_styles(document)

    if client_name:
        info = document.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = info.add_run(f"Клиент: {client_name}")
        run.bold = True

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)

    document.add_paragraph()

    _add_content(document, content)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream


def make_docx_filename(title: str = "document") -> str:
    cleaned = re.sub(r"[^a-zA-Zа-яА-Я0-9_-]+", "_", title.strip())
    cleaned = cleaned.strip("_") or "document"

    return f"{cleaned[:80]}.docx"


def _setup_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)


def _setup_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True


def _add_content(document: Document, content: str) -> None:
    lines = content.splitlines()

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            document.add_paragraph()
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line.replace("## ", "", 1).strip())
            run.bold = True
            run.font.size = Pt(13)
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line.replace("### ", "", 1).strip())
            run.bold = True
            run.font.size = Pt(12)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.add_run("• " + line.replace("- ", "", 1).strip())
            _format_body_paragraph(paragraph)
            continue

        paragraph = document.add_paragraph()
        paragraph.add_run(line)
        _format_body_paragraph(paragraph)


def _format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)