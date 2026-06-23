from pathlib import Path
from uuid import uuid4

import docx2txt
from docx import Document
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from sqlalchemy import text

from app.db.session import SessionLocal


PROJECT_ROOT = Path.cwd().resolve()

UPLOAD_ROOT = PROJECT_ROOT / "uploads"
CASE_UPLOAD_ROOT = UPLOAD_ROOT / "cases"
TEMP_UPLOAD_ROOT = UPLOAD_ROOT / "temp"

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def ensure_document_tables() -> None:
    with SessionLocal() as session:
        session.execute(text("""
            CREATE TABLE IF NOT EXISTS documents (
                id SERIAL PRIMARY KEY,
                case_id INTEGER,
                filename TEXT NOT NULL,
                original_filename TEXT NOT NULL,
                file_path TEXT NOT NULL,
                document_type TEXT,
                extracted_text TEXT,
                summary TEXT,
                created_at TIMESTAMP NOT NULL DEFAULT NOW()
            )
        """))

        session.commit()


def save_case_file(case_id: int, uploaded_file: FileStorage) -> dict:
    ensure_document_tables()

    if not uploaded_file or not uploaded_file.filename:
        raise ValueError("Файл не выбран.")

    original_filename = uploaded_file.filename
    extension = _validate_extension(original_filename)

    CASE_UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)

    safe_name = secure_filename(original_filename) or f"document{extension}"
    stored_filename = f"{uuid4().hex}_{safe_name}"
    file_path = (CASE_UPLOAD_ROOT / stored_filename).resolve()

    uploaded_file.save(file_path)

    try:
        extracted_text = extract_text_from_file(file_path)
    except Exception:
        extracted_text = ""

    with SessionLocal() as session:
        row = session.execute(text("""
            INSERT INTO documents (
                case_id,
                filename,
                original_filename,
                file_path,
                document_type,
                extracted_text
            )
            VALUES (
                :case_id,
                :filename,
                :original_filename,
                :file_path,
                :document_type,
                :extracted_text
            )
            RETURNING
                id,
                case_id,
                filename,
                original_filename,
                file_path,
                document_type,
                extracted_text,
                created_at
        """), {
            "case_id": case_id,
            "filename": stored_filename,
            "original_filename": original_filename,
            "file_path": str(file_path),
            "document_type": extension.replace(".", ""),
            "extracted_text": extracted_text,
        }).mappings().fetchone()

        session.commit()

    return dict(row)


def get_case_documents(case_id: int) -> list[dict]:
    ensure_document_tables()

    with SessionLocal() as session:
        rows = session.execute(text("""
            SELECT
                id,
                case_id,
                filename,
                original_filename,
                file_path,
                document_type,
                created_at
            FROM documents
            WHERE case_id = :case_id
            ORDER BY id DESC
        """), {"case_id": case_id}).mappings().fetchall()

    return [dict(row) for row in rows]


def get_document_file(document_id: int) -> dict:
    ensure_document_tables()

    with SessionLocal() as session:
        row = session.execute(text("""
            SELECT
                id,
                case_id,
                filename,
                original_filename,
                file_path,
                document_type,
                created_at
            FROM documents
            WHERE id = :id
            LIMIT 1
        """), {"id": document_id}).mappings().fetchone()

    if not row:
        raise ValueError("Файл не найден.")

    result = dict(row)
    path = Path(result["file_path"])

    if not path.is_absolute():
        path = (PROJECT_ROOT / path).resolve()

    if not path.exists():
        raise ValueError("Файл отсутствует на диске. Возможно, Render перезапустил сервис и временное файловое хранилище было очищено.")

    result["path"] = str(path)
    return result


def save_temp_file_and_extract(uploaded_file: FileStorage) -> dict:
    if not uploaded_file or not uploaded_file.filename:
        raise ValueError("Файл не выбран.")

    original_filename = uploaded_file.filename
    extension = _validate_extension(original_filename)

    TEMP_UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)

    safe_name = secure_filename(original_filename) or f"document{extension}"
    stored_filename = f"{uuid4().hex}_{safe_name}"
    file_path = (TEMP_UPLOAD_ROOT / stored_filename).resolve()

    uploaded_file.save(file_path)

    extracted_text = extract_text_from_file(file_path)

    if not extracted_text.strip():
        raise ValueError(
            "Не удалось извлечь текст из файла. "
            "Если это PDF-скан, нужен OCR. Если это DOC, сохраните файл как DOCX."
        )

    return {
        "filename": stored_filename,
        "original_filename": original_filename,
        "file_path": str(file_path),
        "extension": extension,
        "text": extracted_text,
    }


def extract_text_from_file(file_path: str | Path) -> str:
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore").strip()

    if extension == ".docx":
        return _extract_docx_text(path)

    if extension == ".pdf":
        return _extract_pdf_text(path)

    raise ValueError("Поддерживаются только PDF, DOCX и TXT.")


def _extract_docx_text(path: Path) -> str:
    parts = []

    try:
        document = Document(str(path))

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = []

                for cell in row.cells:
                    text = cell.text.strip()

                    if text:
                        cells.append(text)

                if cells:
                    parts.append(" | ".join(cells))
    except Exception:
        pass

    text = "\n\n".join(parts).strip()

    if text:
        return text

    try:
        return (docx2txt.process(str(path)) or "").strip()
    except Exception:
        return ""


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []

    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            pages.append(text)

    return "\n\n".join(pages).strip()


def _validate_extension(filename: str) -> str:
    extension = Path(filename).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Поддерживаются только файлы PDF, DOCX и TXT.")

    return extension